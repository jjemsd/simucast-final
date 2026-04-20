# ============================================================================
# services/report_service.py
# ============================================================================
# Generates thesis-ready reports (PDF + DOCX) from a project's timeline.
#
# Design decisions:
#   - Stats and Tests results aren't auto-persisted — the user clicks
#     "Save to report" on the specific results they want included. Those
#     become Steps with step_type="report_item".
#   - Models, cleaning ops, scenarios, and synthetic generations are already
#     persisted as Steps, so the report just reads them straight from the DB.
#   - We use ReportLab for PDF (pure Python, Windows-friendly) and
#     python-docx for DOCX. Neither needs system dependencies.
#   - One gather_report_data() function builds a structured dict. Separate
#     generate_pdf() and generate_docx() both consume that dict, so the
#     content stays identical across formats.
# ============================================================================

import io
import json
from datetime import datetime

from database import db
from models import Dataset, Step

# PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, KeepTogether,
)

# DOCX generation
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ============================================================================
# Orange brand color — matches the UI (#EA580C)
# ============================================================================

BRAND_COLOR = colors.HexColor("#EA580C")
BRAND_COLOR_DOCX = RGBColor(0xEA, 0x58, 0x0C)
MUTED = colors.HexColor("#6B7280")


# ============================================================================
# Save / list / delete report items
# ============================================================================

def save_report_item(project, kind, title, data):
    """
    Store a stats or test result as a Step with step_type="report_item".

    kind: "descriptives" | "frequencies" | "normality" |
          "t_test" | "anova" | "correlation" | "chi_square"
    title: human-readable summary like "Pearson correlation: age vs income"
    data: the full result dict from the service that computed it
    """
    next_order = len(project.steps)
    step = Step(
        project_id=project.id,
        step_type="report_item",
        title=f"Saved to report: {title}",
        order_index=next_order,
        details=json.dumps({
            "kind": kind,
            "title": title,
            "data": data,
        }),
    )
    db.session.add(step)
    db.session.commit()
    return step


def list_report_items(project):
    """Returns all saved report items for a project, in chronological order."""
    steps = Step.query.filter_by(
        project_id=project.id,
        step_type="report_item",
        reverted=False,
    ).order_by(Step.order_index.asc()).all()

    return [{
        "id": s.id,
        "title": s.title,
        "order_index": s.order_index,
        "details": json.loads(s.details),
        "created_at": s.created_at.isoformat(),
    } for s in steps]


def delete_report_item(step_id, user_id):
    """Delete a saved report item. Verifies the user owns it."""
    step = Step.query.get(step_id)
    if not step or step.step_type != "report_item":
        raise ValueError("Item not found")
    if step.project.user_id != user_id:
        raise ValueError("Item not found")
    db.session.delete(step)
    db.session.commit()


# ============================================================================
# Gather all report data
# ============================================================================

def gather_report_data(project):
    """
    Read all relevant Steps + datasets for this project and structure them
    into a dict that the PDF/DOCX generators can iterate over.
    """
    # Current dataset = the most recent non-reverted one
    current_dataset = (
        Dataset.query
        .filter_by(project_id=project.id)
        .order_by(Dataset.id.desc())
        .first()
    )

    # All steps grouped by type (active only — ignore reverted)
    all_steps = (
        Step.query
        .filter_by(project_id=project.id, reverted=False)
        .order_by(Step.order_index.asc())
        .all()
    )

    cleaning_steps = []
    expand_steps = []
    synthetic_steps = []
    saved_items = []
    models = []
    scenarios = []

    for s in all_steps:
        details = json.loads(s.details) if s.details else {}
        bucket = {
            "id": s.id,
            "title": s.title,
            "type": s.step_type,
            "details": details,
            "created_at": s.created_at.isoformat(),
        }

        if s.step_type in ("clean_missing", "clean_outliers", "clean_columns", "clean_duplicates"):
            cleaning_steps.append(bucket)
        elif s.step_type in ("expand_math", "expand_interaction", "expand_bins", "expand_ai"):
            expand_steps.append(bucket)
        elif s.step_type == "synthetic_generated":
            synthetic_steps.append(bucket)
        elif s.step_type == "model_trained":
            models.append(bucket)
        elif s.step_type == "scenario":
            scenarios.append(bucket)
        elif s.step_type == "report_item":
            saved_items.append(bucket)

    return {
        "project": {
            "name": project.name,
            "description": project.description or "",
            "created_at": project.created_at.isoformat(),
            "user_email": project.user.email if project.user else "Unknown",
        },
        "generated_at": datetime.utcnow().isoformat(),
        "current_dataset": {
            "filename": current_dataset.original_filename if current_dataset else None,
            "rows": current_dataset.row_count if current_dataset else 0,
            "cols": current_dataset.column_count if current_dataset else 0,
            "columns_info": json.loads(current_dataset.columns_info) if current_dataset and current_dataset.columns_info else {},
        } if current_dataset else None,
        "cleaning_steps": cleaning_steps,
        "expand_steps": expand_steps,
        "synthetic_steps": synthetic_steps,
        "saved_items": saved_items,
        "models": models,
        "scenarios": scenarios,
    }


# ============================================================================
# PDF generation (ReportLab)
# ============================================================================

def generate_pdf(project):
    """Generate a PDF report. Returns bytes."""
    data = gather_report_data(project)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        rightMargin=0.75*inch, leftMargin=0.75*inch,
        topMargin=0.75*inch, bottomMargin=0.75*inch,
        title=f"{data['project']['name']} - SimuCast Report",
    )

    story = []
    styles = _build_pdf_styles()

    # --- Cover page ---
    story += _pdf_cover(data, styles)
    story.append(PageBreak())

    # --- Dataset overview ---
    if data["current_dataset"]:
        story += _pdf_dataset_section(data, styles)

    # --- Cleaning log ---
    if data["cleaning_steps"] or data["expand_steps"] or data["synthetic_steps"]:
        story += _pdf_cleaning_section(data, styles)

    # --- Saved stats / test results ---
    if data["saved_items"]:
        story += _pdf_saved_items_section(data, styles)

    # --- Trained models ---
    if data["models"]:
        story += _pdf_models_section(data, styles)

    # --- Saved scenarios ---
    if data["scenarios"]:
        story += _pdf_scenarios_section(data, styles)

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def _build_pdf_styles():
    """Custom paragraph styles for the PDF."""
    styles = getSampleStyleSheet()

    styles.add(ParagraphStyle(
        name="BrandTitle",
        parent=styles["Title"],
        textColor=BRAND_COLOR,
        fontSize=26,
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="BrandH1",
        parent=styles["Heading1"],
        textColor=BRAND_COLOR,
        fontSize=16,
        spaceBefore=12,
        spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name="BrandH2",
        parent=styles["Heading2"],
        textColor=colors.HexColor("#9A3412"),
        fontSize=12,
        spaceBefore=8,
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="Muted",
        parent=styles["Normal"],
        textColor=MUTED,
        fontSize=9,
    ))
    styles.add(ParagraphStyle(
        name="Body",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=6,
        leading=14,
    ))
    styles.add(ParagraphStyle(
        name="Mono",
        parent=styles["Code"],
        fontSize=9,
        backColor=colors.HexColor("#F3F4F6"),
    ))
    return styles


def _pdf_cover(data, styles):
    """Cover page elements."""
    p = data["project"]
    story = []

    story.append(Spacer(1, 1.5*inch))
    story.append(Paragraph(p["name"], styles["BrandTitle"]))
    story.append(Paragraph("SimuCast Analysis Report", styles["Heading2"]))
    story.append(Spacer(1, 0.4*inch))

    if p["description"]:
        story.append(Paragraph(p["description"], styles["Body"]))
        story.append(Spacer(1, 0.2*inch))

    # Metadata table
    meta = [
        ["Generated", datetime.fromisoformat(data["generated_at"]).strftime("%B %d, %Y · %H:%M UTC")],
        ["Author",    p["user_email"]],
        ["Project created", datetime.fromisoformat(p["created_at"]).strftime("%B %d, %Y")],
    ]
    if data["current_dataset"]:
        ds = data["current_dataset"]
        meta.append(["Dataset", f"{ds['filename']} · {ds['rows']:,} rows · {ds['cols']} columns"])

    t = Table(meta, colWidths=[1.5*inch, 4.5*inch])
    t.setStyle(TableStyle([
        ("TEXTCOLOR", (0, 0), (0, -1), MUTED),
        ("TEXTCOLOR", (1, 0), (1, -1), colors.black),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(t)

    return story


def _pdf_dataset_section(data, styles):
    ds = data["current_dataset"]
    story = [Paragraph("1. Dataset overview", styles["BrandH1"])]

    story.append(Paragraph(
        f"<b>{ds['filename']}</b> — {ds['rows']:,} rows across {ds['cols']} columns.",
        styles["Body"],
    ))
    story.append(Spacer(1, 0.1*inch))

    # Columns table
    col_rows = [["Column", "Type"]]
    for col, dtype in ds["columns_info"].items():
        col_rows.append([col, dtype])

    t = Table(col_rows, colWidths=[3*inch, 2*inch])
    t.setStyle(_standard_table_style())
    story.append(t)
    story.append(Spacer(1, 0.2*inch))

    return story


def _pdf_cleaning_section(data, styles):
    story = [Paragraph("2. Data preparation log", styles["BrandH1"])]

    all_prep = data["synthetic_steps"] + data["cleaning_steps"] + data["expand_steps"]
    if not all_prep:
        return []

    rows = [["#", "Action", "Type"]]
    for i, step in enumerate(all_prep, 1):
        kind_label = _STEP_TYPE_LABELS.get(step["type"], step["type"])
        rows.append([str(i), step["title"], kind_label])

    t = Table(rows, colWidths=[0.3*inch, 4.2*inch, 1.5*inch])
    t.setStyle(_standard_table_style())
    story.append(t)
    story.append(Spacer(1, 0.2*inch))

    return story


_STEP_TYPE_LABELS = {
    "clean_missing":        "Missing values",
    "clean_outliers":       "Outliers",
    "clean_columns":        "Columns",
    "clean_duplicates":     "Duplicates",
    "expand_math":          "Math transform",
    "expand_interaction":   "Interaction",
    "expand_bins":          "Binning",
    "expand_ai":            "AI feature",
    "synthetic_generated":  "Synthetic data",
}


def _pdf_saved_items_section(data, styles):
    story = [Paragraph("3. Statistical analyses", styles["BrandH1"])]

    for idx, item in enumerate(data["saved_items"], 1):
        details = item["details"]
        kind = details.get("kind", "")
        title = details.get("title", "Untitled")
        result_data = details.get("data", {})

        story.append(Paragraph(f"3.{idx} {title}", styles["BrandH2"]))
        story += _pdf_render_item(kind, result_data, styles)
        story.append(Spacer(1, 0.15*inch))

    return story


def _pdf_render_item(kind, data, styles):
    """Render one saved result into PDF flowables."""
    if kind == "descriptives":
        return _pdf_descriptives(data, styles)
    if kind == "frequencies":
        return _pdf_frequencies(data, styles)
    if kind == "normality":
        return _pdf_normality(data, styles)
    if kind == "t_test":
        return _pdf_t_test(data, styles)
    if kind == "anova":
        return _pdf_anova(data, styles)
    if kind == "correlation":
        return _pdf_correlation(data, styles)
    if kind == "chi_square":
        return _pdf_chi_square(data, styles)
    return [Paragraph(f"Unknown result type: {kind}", styles["Muted"])]


def _pdf_descriptives(data, styles):
    """data shape: { column_name: {mean, median, std, ...} }"""
    rows = [["Column", "N", "Mean", "Median", "SD", "Min", "Max", "Skew"]]
    for col, s in data.items():
        if "error" in s:
            continue
        rows.append([
            col,
            str(s.get("count", "")),
            _fmt(s.get("mean")),
            _fmt(s.get("median")),
            _fmt(s.get("std")),
            _fmt(s.get("min")),
            _fmt(s.get("max")),
            _fmt(s.get("skewness")),
        ])
    t = Table(rows, repeatRows=1)
    t.setStyle(_standard_table_style())
    return [t]


def _pdf_frequencies(data, styles):
    rows = [["Value", "Count", "Percent"]]
    for f in data.get("frequencies", [])[:20]:   # Cap at 20 for PDF brevity
        rows.append([str(f["value"]), str(f["count"]), f"{f['percent']}%"])

    t = Table(rows, colWidths=[3*inch, 1*inch, 1*inch])
    t.setStyle(_standard_table_style())
    return [
        Paragraph(
            f"Column: <b>{data.get('column', '')}</b> · {data.get('unique_count', 0)} unique values · {data.get('total', 0)} total rows",
            styles["Body"],
        ),
        t,
    ]


def _pdf_normality(data, styles):
    is_normal = data.get("is_normal")
    verdict = "Normal" if is_normal else "Not normal"
    color_note = "darkgreen" if is_normal else "darkorange"

    return [
        Paragraph(
            f"Column: <b>{data.get('column', '')}</b> · n = {data.get('n', 0)}",
            styles["Body"],
        ),
        Paragraph(
            f"Verdict: <font color='{color_note}'><b>{verdict}</b></font> "
            f"(Shapiro-Wilk W = {data.get('shapiro_statistic')}, p = {data.get('shapiro_p')})",
            styles["Body"],
        ),
        Paragraph(
            f"Skewness: {data.get('skewness')} · Kurtosis: {data.get('kurtosis')}",
            styles["Muted"],
        ),
    ]


def _pdf_t_test(data, styles):
    test = data.get("test", "t-test")
    sig = data.get("significant")
    sig_text = "SIGNIFICANT" if sig else "not significant"
    sig_color = "darkgreen" if sig else "gray"

    rows = [
        ["Result", f"{sig_text} at α = 0.05"],
        ["p-value", _fmt(data.get("p_value"), 6)],
        ["t statistic", _fmt(data.get("t_statistic"))],
    ]
    if "degrees_of_freedom" in data:
        rows.append(["Degrees of freedom", str(data["degrees_of_freedom"])])
    if "cohens_d" in data:
        rows.append(["Cohen's d", _fmt(data.get("cohens_d"))])

    # Test-specific extras
    if test == "independent_t":
        for g in ["group_1", "group_2"]:
            gd = data.get(g, {})
            rows.append([f"{gd.get('name', g)} (n={gd.get('n')})", f"mean {gd.get('mean')}, SD {gd.get('std')}"])
        rows.append(["Mean difference", _fmt(data.get("mean_difference"))])
    elif test == "paired_t":
        rows.append([f"Mean of {data.get('column_a')}", _fmt(data.get("mean_a"))])
        rows.append([f"Mean of {data.get('column_b')}", _fmt(data.get("mean_b"))])
        rows.append(["Mean difference", _fmt(data.get("mean_difference"))])
    elif test == "one_sample_t":
        rows.append(["Test value", _fmt(data.get("test_value"))])
        rows.append(["Sample mean", _fmt(data.get("sample_mean"))])

    t = Table(rows, colWidths=[2*inch, 4*inch])
    t.setStyle(_kv_table_style())
    return [
        Paragraph(
            f"<b>Test:</b> {test.replace('_', ' ')} · "
            f"<font color='{sig_color}'><b>{sig_text}</b></font>",
            styles["Body"],
        ),
        t,
    ]


def _pdf_anova(data, styles):
    rows = [
        ["F statistic", _fmt(data.get("f_statistic"))],
        ["p value", _fmt(data.get("p_value"), 6)],
        ["df between / within", f"{data.get('df_between')} / {data.get('df_within')}"],
        ["η² (effect size)", _fmt(data.get("eta_squared"))],
        ["Total n", str(data.get("n", ""))],
        ["Groups", str(data.get("num_groups", ""))],
    ]
    kv = Table(rows, colWidths=[2*inch, 4*inch])
    kv.setStyle(_kv_table_style())

    # Group breakdown table
    group_rows = [["Group", "n", "Mean", "SD"]]
    for g in data.get("groups", []):
        group_rows.append([str(g["name"]), str(g["n"]), _fmt(g["mean"]), _fmt(g["std"])])
    group_t = Table(group_rows, colWidths=[2*inch, 0.8*inch, 1.2*inch, 1.2*inch])
    group_t.setStyle(_standard_table_style())

    return [
        Paragraph(
            f"<b>Comparing {data.get('numeric_column')} by {data.get('group_column')}</b>",
            styles["Body"],
        ),
        kv, Spacer(1, 0.1*inch),
        Paragraph("Group means:", styles["Body"]),
        group_t,
    ]


def _pdf_correlation(data, styles):
    r = data.get("r", 0)
    sig = data.get("significant")
    direction = data.get("direction", "")
    strength = data.get("strength", "")

    rows = [
        ["Method", data.get("method", "").capitalize()],
        ["r (correlation)", _fmt(r)],
        ["r² (shared variance)", _fmt(data.get("r_squared"))],
        ["p value", _fmt(data.get("p_value"), 6)],
        ["n", str(data.get("n", ""))],
        ["Strength", f"{strength} {direction}"],
        ["Significance", "significant at α=0.05" if sig else "not significant"],
    ]
    t = Table(rows, colWidths=[2*inch, 4*inch])
    t.setStyle(_kv_table_style())
    return [
        Paragraph(
            f"<b>Correlating {data.get('column_a')} with {data.get('column_b')}</b>",
            styles["Body"],
        ),
        t,
    ]


def _pdf_chi_square(data, styles):
    rows = [
        ["χ² statistic", _fmt(data.get("chi2"))],
        ["Degrees of freedom", str(data.get("degrees_of_freedom", ""))],
        ["p value", _fmt(data.get("p_value"), 6)],
        ["Cramér's V", _fmt(data.get("cramers_v"))],
        ["n", str(data.get("n", ""))],
    ]
    kv = Table(rows, colWidths=[2*inch, 4*inch])
    kv.setStyle(_kv_table_style())

    # Contingency table
    ct = data.get("contingency_table", {})
    rlabels = ct.get("rows", [])
    clabels = ct.get("columns", [])
    values = ct.get("values", [])

    cross_rows = [[""] + clabels]
    for i, r in enumerate(rlabels):
        cross_rows.append([r] + [str(v) for v in values[i]])

    ct_table = Table(cross_rows)
    ct_table.setStyle(_standard_table_style())

    return [
        Paragraph(
            f"<b>{data.get('column_a')} × {data.get('column_b')}</b>",
            styles["Body"],
        ),
        kv, Spacer(1, 0.1*inch),
        Paragraph("Contingency table:", styles["Body"]),
        ct_table,
    ]


def _pdf_models_section(data, styles):
    story = [Paragraph("4. Predictive models", styles["BrandH1"])]

    for idx, m in enumerate(data["models"], 1):
        details = m["details"]
        metrics = details.get("metrics", {})
        story.append(Paragraph(f"4.{idx} {m['title']}", styles["BrandH2"]))

        # Model metadata
        meta_rows = [
            ["Type",     details.get("model_type", "")],
            ["Target",   details.get("target", "")],
            ["Features", ", ".join(details.get("features_original", []))],
            ["Rows used", str(details.get("n_rows_used", ""))],
        ]
        meta_t = Table(meta_rows, colWidths=[1.5*inch, 4.5*inch])
        meta_t.setStyle(_kv_table_style())
        story.append(meta_t)
        story.append(Spacer(1, 0.1*inch))

        # Metrics — show the right set depending on regression vs classification
        metric_rows = [["Metric", "Value"]]
        for key in ["r2_test", "r2_train", "rmse_test", "mae_test",
                    "accuracy", "precision", "recall", "f1", "roc_auc"]:
            if key in metrics:
                metric_rows.append([key.replace("_", " "), _fmt(metrics[key])])

        m_t = Table(metric_rows, colWidths=[2*inch, 2*inch])
        m_t.setStyle(_standard_table_style())
        story.append(m_t)

        # Top features
        feature_list = metrics.get("feature_importances") or metrics.get("coefficients")
        if feature_list:
            story.append(Spacer(1, 0.1*inch))
            story.append(Paragraph("Top features:", styles["Body"]))
            sorted_feats = sorted(feature_list, key=lambda x: abs(x["value"]), reverse=True)[:10]
            feat_rows = [["Feature", "Value"]]
            for f in sorted_feats:
                feat_rows.append([f["feature"], _fmt(f["value"])])
            feat_t = Table(feat_rows, colWidths=[3.5*inch, 2*inch])
            feat_t.setStyle(_standard_table_style())
            story.append(feat_t)

        story.append(Spacer(1, 0.2*inch))

    return story


def _pdf_scenarios_section(data, styles):
    story = [Paragraph("5. What-if scenarios", styles["BrandH1"])]

    # Group scenarios by model
    by_model = {}
    for s in data["scenarios"]:
        details = s["details"]
        model_id = details.get("model_step_id", 0)
        by_model.setdefault(model_id, []).append(details)

    for model_id, scenarios in by_model.items():
        if not scenarios:
            continue
        story.append(Paragraph(f"Scenarios for model #{model_id}", styles["BrandH2"]))

        for sc in scenarios:
            pred = sc.get("prediction", {})
            pred_str = str(pred.get("value", "—"))
            story.append(Paragraph(
                f"<b>{sc.get('name', 'Unnamed')}</b> — predicts: {pred_str}",
                styles["Body"],
            ))
            inputs = sc.get("inputs", {})
            input_str = " · ".join(f"{k}={v}" for k, v in inputs.items())
            story.append(Paragraph(input_str, styles["Muted"]))
            story.append(Spacer(1, 0.08*inch))

    return story


# --- Table styling helpers ---

def _standard_table_style():
    """Header row in brand color, subtle grid lines."""
    return TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#FFEDD5")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#9A3412")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#E5E7EB")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#FAFAFA")]),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ])


def _kv_table_style():
    """For key/value tables — no header row, muted labels."""
    return TableStyle([
        ("TEXTCOLOR", (0, 0), (0, -1), MUTED),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
        ("LINEBELOW", (0, 0), (-1, -2), 0.25, colors.HexColor("#F3F4F6")),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ])


def _fmt(v, decimals=4):
    """Format a number with sensible precision. Returns '—' for None/NaN."""
    if v is None:
        return "—"
    try:
        return f"{float(v):.{decimals}f}".rstrip("0").rstrip(".") or "0"
    except (TypeError, ValueError):
        return str(v)


# ============================================================================
# DOCX generation (python-docx)
# ============================================================================

def generate_docx(project):
    """Generate a Word doc. Returns bytes."""
    data = gather_report_data(project)
    doc = Document()

    # --- Cover ---
    _docx_cover(doc, data)
    doc.add_page_break()

    # --- Dataset ---
    if data["current_dataset"]:
        _docx_dataset_section(doc, data)

    # --- Cleaning log ---
    if data["cleaning_steps"] or data["expand_steps"] or data["synthetic_steps"]:
        _docx_cleaning_section(doc, data)

    # --- Saved items ---
    if data["saved_items"]:
        _docx_saved_items_section(doc, data)

    # --- Models ---
    if data["models"]:
        _docx_models_section(doc, data)

    # --- Scenarios ---
    if data["scenarios"]:
        _docx_scenarios_section(doc, data)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _docx_cover(doc, data):
    p = data["project"]
    title = doc.add_heading(p["name"], level=0)
    for run in title.runs:
        run.font.color.rgb = BRAND_COLOR_DOCX

    subtitle = doc.add_paragraph("SimuCast Analysis Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in subtitle.runs:
        run.font.size = Pt(14)

    if p["description"]:
        doc.add_paragraph(p["description"])

    doc.add_paragraph()  # spacer

    # Metadata
    gen_date = datetime.fromisoformat(data["generated_at"]).strftime("%B %d, %Y · %H:%M UTC")
    meta = [
        ("Generated", gen_date),
        ("Author", p["user_email"]),
        ("Project created", datetime.fromisoformat(p["created_at"]).strftime("%B %d, %Y")),
    ]
    if data["current_dataset"]:
        ds = data["current_dataset"]
        meta.append(("Dataset", f"{ds['filename']} · {ds['rows']:,} rows · {ds['cols']} columns"))

    table = doc.add_table(rows=len(meta), cols=2)
    for i, (k, v) in enumerate(meta):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = str(v)
        # Label column grey
        for run in table.cell(i, 0).paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def _docx_section_heading(doc, text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = BRAND_COLOR_DOCX


def _docx_dataset_section(doc, data):
    ds = data["current_dataset"]
    _docx_section_heading(doc, "1. Dataset overview")

    doc.add_paragraph(f"{ds['filename']} — {ds['rows']:,} rows across {ds['cols']} columns.")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 2"
    hdr = table.rows[0].cells
    hdr[0].text = "Column"
    hdr[1].text = "Type"
    for col, dtype in ds["columns_info"].items():
        row = table.add_row().cells
        row[0].text = col
        row[1].text = dtype


def _docx_cleaning_section(doc, data):
    _docx_section_heading(doc, "2. Data preparation log")
    all_prep = data["synthetic_steps"] + data["cleaning_steps"] + data["expand_steps"]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 2"
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Action"
    hdr[2].text = "Type"
    for i, step in enumerate(all_prep, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = step["title"]
        row[2].text = _STEP_TYPE_LABELS.get(step["type"], step["type"])


def _docx_saved_items_section(doc, data):
    _docx_section_heading(doc, "3. Statistical analyses")

    for idx, item in enumerate(data["saved_items"], 1):
        details = item["details"]
        kind = details.get("kind", "")
        title = details.get("title", "Untitled")
        result_data = details.get("data", {})

        h = doc.add_heading(f"3.{idx} {title}", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x9A, 0x34, 0x12)

        _docx_render_item(doc, kind, result_data)


def _docx_render_item(doc, kind, data):
    """Format one result into Word paragraphs + tables."""

    if kind == "descriptives":
        table = doc.add_table(rows=1, cols=8)
        table.style = "Light Grid Accent 2"
        for i, h in enumerate(["Column", "N", "Mean", "Median", "SD", "Min", "Max", "Skew"]):
            table.rows[0].cells[i].text = h
        for col, s in data.items():
            if "error" in s:
                continue
            cells = table.add_row().cells
            cells[0].text = col
            cells[1].text = str(s.get("count", ""))
            cells[2].text = _fmt(s.get("mean"))
            cells[3].text = _fmt(s.get("median"))
            cells[4].text = _fmt(s.get("std"))
            cells[5].text = _fmt(s.get("min"))
            cells[6].text = _fmt(s.get("max"))
            cells[7].text = _fmt(s.get("skewness"))

    elif kind == "frequencies":
        doc.add_paragraph(
            f"Column: {data.get('column', '')} · {data.get('unique_count', 0)} unique · {data.get('total', 0)} rows"
        )
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 2"
        for i, h in enumerate(["Value", "Count", "Percent"]):
            table.rows[0].cells[i].text = h
        for f in data.get("frequencies", [])[:20]:
            cells = table.add_row().cells
            cells[0].text = str(f["value"])
            cells[1].text = str(f["count"])
            cells[2].text = f"{f['percent']}%"

    elif kind == "normality":
        verdict = "Normal" if data.get("is_normal") else "Not normal"
        doc.add_paragraph(f"Column: {data.get('column', '')} · n = {data.get('n', 0)}")
        doc.add_paragraph(
            f"Verdict: {verdict} (Shapiro-Wilk W = {data.get('shapiro_statistic')}, p = {data.get('shapiro_p')})"
        )
        doc.add_paragraph(
            f"Skewness: {data.get('skewness')} · Kurtosis: {data.get('kurtosis')}"
        )

    elif kind == "t_test":
        test = data.get("test", "t-test")
        sig = "significant" if data.get("significant") else "not significant"
        doc.add_paragraph(f"Test: {test.replace('_', ' ')} — {sig} at α = 0.05")
        _docx_kv_table(doc, [
            ("p-value", _fmt(data.get("p_value"), 6)),
            ("t statistic", _fmt(data.get("t_statistic"))),
            ("Cohen's d", _fmt(data.get("cohens_d"))),
            ("Degrees of freedom", str(data.get("degrees_of_freedom", "—"))),
        ])

    elif kind == "anova":
        doc.add_paragraph(
            f"Comparing {data.get('numeric_column')} across groups of {data.get('group_column')}"
        )
        _docx_kv_table(doc, [
            ("F statistic", _fmt(data.get("f_statistic"))),
            ("p value", _fmt(data.get("p_value"), 6)),
            ("η² (effect size)", _fmt(data.get("eta_squared"))),
            ("df between / within", f"{data.get('df_between')} / {data.get('df_within')}"),
            ("Total n", str(data.get("n", ""))),
        ])
        doc.add_paragraph("Group means:")
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 2"
        for i, h in enumerate(["Group", "n", "Mean", "SD"]):
            table.rows[0].cells[i].text = h
        for g in data.get("groups", []):
            cells = table.add_row().cells
            cells[0].text = str(g["name"])
            cells[1].text = str(g["n"])
            cells[2].text = _fmt(g["mean"])
            cells[3].text = _fmt(g["std"])

    elif kind == "correlation":
        doc.add_paragraph(
            f"Correlating {data.get('column_a')} with {data.get('column_b')}"
        )
        _docx_kv_table(doc, [
            ("Method", (data.get("method") or "").capitalize()),
            ("r (correlation)", _fmt(data.get("r"))),
            ("r² (shared variance)", _fmt(data.get("r_squared"))),
            ("p value", _fmt(data.get("p_value"), 6)),
            ("n", str(data.get("n", ""))),
            ("Strength", f"{data.get('strength', '')} {data.get('direction', '')}"),
        ])

    elif kind == "chi_square":
        doc.add_paragraph(f"{data.get('column_a')} × {data.get('column_b')}")
        _docx_kv_table(doc, [
            ("χ² statistic", _fmt(data.get("chi2"))),
            ("p value", _fmt(data.get("p_value"), 6)),
            ("Cramér's V", _fmt(data.get("cramers_v"))),
            ("Degrees of freedom", str(data.get("degrees_of_freedom", ""))),
            ("n", str(data.get("n", ""))),
        ])
        # Contingency table
        ct = data.get("contingency_table", {})
        rlabels = ct.get("rows", [])
        clabels = ct.get("columns", [])
        values = ct.get("values", [])
        if rlabels and clabels:
            doc.add_paragraph("Contingency table:")
            table = doc.add_table(rows=1, cols=len(clabels) + 1)
            table.style = "Light Grid Accent 2"
            table.rows[0].cells[0].text = ""
            for i, c in enumerate(clabels, 1):
                table.rows[0].cells[i].text = str(c)
            for i, r in enumerate(rlabels):
                cells = table.add_row().cells
                cells[0].text = str(r)
                for j, v in enumerate(values[i], 1):
                    cells[j].text = str(v)

    doc.add_paragraph()  # spacer


def _docx_kv_table(doc, rows):
    """Simple 2-column key/value table."""
    table = doc.add_table(rows=len(rows), cols=2)
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = str(v)
        for run in table.cell(i, 0).paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def _docx_models_section(doc, data):
    _docx_section_heading(doc, "4. Predictive models")

    for idx, m in enumerate(data["models"], 1):
        details = m["details"]
        metrics = details.get("metrics", {})

        h = doc.add_heading(f"4.{idx} {m['title']}", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x9A, 0x34, 0x12)

        _docx_kv_table(doc, [
            ("Type", details.get("model_type", "")),
            ("Target", details.get("target", "")),
            ("Features", ", ".join(details.get("features_original", []))),
            ("Rows used", str(details.get("n_rows_used", ""))),
        ])

        doc.add_paragraph("Metrics:")
        metric_rows = []
        for key in ["r2_test", "r2_train", "rmse_test", "mae_test",
                    "accuracy", "precision", "recall", "f1", "roc_auc"]:
            if key in metrics:
                metric_rows.append((key.replace("_", " "), _fmt(metrics[key])))
        _docx_kv_table(doc, metric_rows)

        feature_list = metrics.get("feature_importances") or metrics.get("coefficients")
        if feature_list:
            doc.add_paragraph("Top features:")
            sorted_feats = sorted(feature_list, key=lambda x: abs(x["value"]), reverse=True)[:10]
            _docx_kv_table(doc, [(f["feature"], _fmt(f["value"])) for f in sorted_feats])

        doc.add_paragraph()


def _docx_scenarios_section(doc, data):
    _docx_section_heading(doc, "5. What-if scenarios")

    for s in data["scenarios"]:
        details = s["details"]
        pred = details.get("prediction", {})
        name = details.get("name", "Unnamed")
        pred_val = str(pred.get("value", "—"))

        p = doc.add_paragraph()
        p.add_run(f"{name}").bold = True
        p.add_run(f" — predicts: {pred_val}")

        inputs = details.get("inputs", {})
        input_str = " · ".join(f"{k}={v}" for k, v in inputs.items())
        muted = doc.add_paragraph(input_str)
        for run in muted.runs:
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            run.font.size = Pt(9)
